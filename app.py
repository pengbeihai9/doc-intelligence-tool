import hashlib
import io
import os
import re
import secrets
import sqlite3
import time
from collections import Counter
from pathlib import Path

import fitz
import jieba
import matplotlib.pyplot as plt
import pandas as pd
import plotly.express as px
import pytesseract
import streamlit as st
from docx import Document
from pdf2image import convert_from_bytes
from PIL import Image
from wordcloud import WordCloud


BASE_DIR = Path(__file__).resolve().parent
DATA_DIR = BASE_DIR / "data"
UPLOAD_DIR = DATA_DIR / "uploads"
DB_PATH = DATA_DIR / "app.db"

STOPWORDS = {
    "的", "了", "和", "是", "在", "有", "与", "及", "或", "为", "对", "等", "中",
    "the", "and", "for", "with", "that", "this", "from", "are", "was", "were"
}


# ---------- storage and auth ----------

def init_storage() -> None:
    DATA_DIR.mkdir(exist_ok=True)
    UPLOAD_DIR.mkdir(exist_ok=True)
    with sqlite3.connect(DB_PATH) as conn:
        conn.execute(
            """
            CREATE TABLE IF NOT EXISTS users (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                username TEXT NOT NULL UNIQUE,
                password_hash TEXT NOT NULL,
                salt TEXT NOT NULL,
                created_at TEXT NOT NULL DEFAULT CURRENT_TIMESTAMP
            )
            """
        )
        conn.execute(
            """
            CREATE TABLE IF NOT EXISTS documents (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                user_id INTEGER NOT NULL,
                filename TEXT NOT NULL,
                file_path TEXT NOT NULL,
                category TEXT NOT NULL,
                summary TEXT NOT NULL,
                extracted_text TEXT NOT NULL,
                word_count INTEGER NOT NULL,
                created_at TEXT NOT NULL DEFAULT CURRENT_TIMESTAMP,
                FOREIGN KEY(user_id) REFERENCES users(id)
            )
            """
        )


def hash_password(password: str, salt: str | None = None) -> tuple[str, str]:
    salt = salt or secrets.token_hex(16)
    digest = hashlib.pbkdf2_hmac("sha256", password.encode("utf-8"), salt.encode("utf-8"), 120_000)
    return digest.hex(), salt


def create_user(username: str, password: str) -> tuple[bool, str]:
    username = username.strip()
    if not re.fullmatch(r"[A-Za-z0-9_\u4e00-\u9fff]{2,20}", username):
        return False, "用户名需为 2-20 位中文、字母、数字或下划线。"
    if len(password) < 6:
        return False, "密码至少需要 6 位。"

    password_hash, salt = hash_password(password)
    try:
        with sqlite3.connect(DB_PATH) as conn:
            conn.execute(
                "INSERT INTO users (username, password_hash, salt) VALUES (?, ?, ?)",
                (username, password_hash, salt),
            )
        return True, "注册成功，请登录。"
    except sqlite3.IntegrityError:
        return False, "用户名已存在。"


def authenticate(username: str, password: str) -> dict | None:
    with sqlite3.connect(DB_PATH) as conn:
        conn.row_factory = sqlite3.Row
        user = conn.execute("SELECT * FROM users WHERE username = ?", (username.strip(),)).fetchone()
    if user is None:
        return None
    password_hash, _ = hash_password(password, user["salt"])
    if not secrets.compare_digest(password_hash, user["password_hash"]):
        return None
    return {"id": user["id"], "username": user["username"]}


def save_document(user_id: int, filename: str, file_bytes: bytes, category: str, summary: str, extracted_text: str, word_count: int) -> None:
    user_dir = UPLOAD_DIR / str(user_id)
    user_dir.mkdir(parents=True, exist_ok=True)
    safe_name = re.sub(r"[^A-Za-z0-9_.\-\u4e00-\u9fff]", "_", filename)
    stored_name = f"{int(time.time())}_{safe_name}"
    file_path = user_dir / stored_name
    file_path.write_bytes(file_bytes)

    with sqlite3.connect(DB_PATH) as conn:
        conn.execute(
            """
            INSERT INTO documents
            (user_id, filename, file_path, category, summary, extracted_text, word_count)
            VALUES (?, ?, ?, ?, ?, ?, ?)
            """,
            (user_id, filename, str(file_path), category, summary, extracted_text, word_count),
        )


def list_user_documents(user_id: int) -> list[sqlite3.Row]:
    with sqlite3.connect(DB_PATH) as conn:
        conn.row_factory = sqlite3.Row
        return conn.execute(
            """
            SELECT id, filename, category, summary, word_count, created_at
            FROM documents
            WHERE user_id = ?
            ORDER BY created_at DESC, id DESC
            """,
            (user_id,),
        ).fetchall()


def get_user_document(user_id: int, document_id: int) -> sqlite3.Row | None:
    with sqlite3.connect(DB_PATH) as conn:
        conn.row_factory = sqlite3.Row
        return conn.execute(
            """
            SELECT * FROM documents
            WHERE id = ? AND user_id = ?
            """,
            (document_id, user_id),
        ).fetchone()


# ---------- document processing ----------

def clean_text(text: str) -> str:
    text = text.replace("\x00", " ")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def extract_pdf(file_bytes: bytes) -> str:
    doc = fitz.open(stream=file_bytes, filetype="pdf")
    pages_text = [page.get_text("text") for page in doc]
    text = clean_text("\n".join(pages_text))

    if len(text) >= 80:
        return text

    images = convert_from_bytes(file_bytes, dpi=200)
    ocr_text = []
    for image in images:
        ocr_text.append(pytesseract.image_to_string(image, lang="chi_sim+eng"))
    return clean_text("\n".join(ocr_text))


def extract_docx(file_obj) -> str:
    document = Document(file_obj)
    parts = []

    for paragraph in document.paragraphs:
        if paragraph.text.strip():
            parts.append(paragraph.text)

    for table in document.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                parts.append(row_text)

    return clean_text("\n".join(parts))


def extract_image(file_obj) -> str:
    image = Image.open(file_obj)
    return clean_text(pytesseract.image_to_string(image, lang="chi_sim+eng"))


def split_sentences(text: str) -> list[str]:
    sentences = re.split(r"(?<=[。！？.!?])\s*", text)
    return [sentence.strip() for sentence in sentences if len(sentence.strip()) >= 12]


def tokenize(text: str) -> list[str]:
    words = []
    for word in jieba.cut(text):
        word = word.strip().lower()
        if len(word) < 2:
            continue
        if word in STOPWORDS:
            continue
        if re.fullmatch(r"[\d\W_]+", word):
            continue
        words.append(word)
    return words


def summarize_text(text: str, max_sentences: int = 5) -> str:
    sentences = split_sentences(text)
    if not sentences:
        return text[:500]

    word_freq = Counter(tokenize(text))
    scored = []
    for index, sentence in enumerate(sentences):
        score = sum(word_freq.get(word, 0) for word in tokenize(sentence))
        score = score / max(len(sentence), 1)
        scored.append((score, index, sentence))

    selected = sorted(scored, reverse=True)[:max_sentences]
    selected = sorted(selected, key=lambda item: item[1])
    return "\n".join(item[2] for item in selected)


def classify_document(text: str) -> tuple[str, dict[str, int]]:
    rules = {
        "合同/协议": ["合同", "协议", "甲方", "乙方", "违约", "签订", "条款", "履行"],
        "财务/票据": ["发票", "金额", "税率", "付款", "收款", "银行", "费用", "报销"],
        "简历/人才材料": ["简历", "教育经历", "工作经历", "项目经验", "技能", "求职", "实习"],
        "论文/研究报告": ["摘要", "关键词", "研究", "实验", "模型", "数据", "结论", "参考文献"],
        "通知/公文": ["通知", "公告", "会议", "决定", "要求", "单位", "日期", "请于"],
    }

    scores = {
        category: sum(text.count(keyword) for keyword in keywords)
        for category, keywords in rules.items()
    }
    best_category = max(scores, key=scores.get)
    if scores[best_category] == 0:
        return "其他文档", scores
    return best_category, scores


def build_wordcloud(freq_df: pd.DataFrame):
    frequencies = dict(zip(freq_df["词语"], freq_df["频次"]))
    font_candidates = [
        "C:/Windows/Fonts/simhei.ttf",
        "C:/Windows/Fonts/msyh.ttc",
        "/usr/share/fonts/truetype/wqy/wqy-microhei.ttc",
        "/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc",
    ]
    font_path = next((path for path in font_candidates if Path(path).exists()), None)
    wordcloud = WordCloud(
        font_path=font_path,
        width=900,
        height=420,
        background_color="white",
        max_words=80,
    ).generate_from_frequencies(frequencies)

    figure, axis = plt.subplots(figsize=(10, 4.8))
    axis.imshow(wordcloud, interpolation="bilinear")
    axis.axis("off")
    return figure


def parse_uploaded_file(uploaded_file) -> str:
    suffix = uploaded_file.name.lower().rsplit(".", 1)[-1]
    file_bytes = uploaded_file.getvalue()

    if suffix == "pdf":
        return extract_pdf(file_bytes)
    if suffix == "docx":
        return extract_docx(io.BytesIO(file_bytes))
    if suffix in {"png", "jpg", "jpeg", "bmp", "tif", "tiff"}:
        return extract_image(io.BytesIO(file_bytes))

    raise ValueError("暂不支持该文件格式")


def analyze_text(text: str) -> tuple[str, dict[str, int], str, pd.DataFrame]:
    category, category_scores = classify_document(text)
    summary = summarize_text(text)
    freq = Counter(tokenize(text)).most_common(30)
    freq_df = pd.DataFrame(freq, columns=["词语", "频次"])
    return category, category_scores, summary, freq_df


def render_analysis(category: str, category_scores: dict[str, int], summary: str, freq_df: pd.DataFrame, extracted_text: str) -> None:
    left, right = st.columns([1, 1])

    with left:
        st.subheader("自动分类")
        st.success(category)
        st.caption("分类方式：基于关键词规则打分，可按实际业务继续扩展。")

    with right:
        st.subheader("文档信息")
        st.metric("文本长度", f"{len(extracted_text)} 字符")
        st.metric("关键词数量", len(freq_df))

    st.subheader("文本摘要")
    st.write(summary)

    st.subheader("词频统计")
    if freq_df.empty:
        st.write("暂无可统计词语。")
    else:
        tab_bar, tab_cloud, tab_class = st.tabs(["柱状图", "词云", "分类得分"])

        with tab_bar:
            fig = px.bar(freq_df, x="词语", y="频次", text="频次", title="高频词 Top 30")
            fig.update_layout(xaxis_title="词语", yaxis_title="频次")
            st.plotly_chart(fig, use_container_width=True)

        with tab_cloud:
            try:
                st.pyplot(build_wordcloud(freq_df), clear_figure=True)
            except Exception as exc:
                st.write(f"词云生成失败：{exc}。柱状图仍可正常使用。")

        with tab_class:
            score_df = pd.DataFrame(
                sorted(category_scores.items(), key=lambda item: item[1], reverse=True),
                columns=["类别", "得分"],
            )
            st.dataframe(score_df, use_container_width=True)

    st.subheader("提取的全文")
    st.text_area("全文内容", extracted_text, height=360)


# ---------- Streamlit UI ----------

init_storage()
st.set_page_config(page_title="文档智能处理系统", page_icon="📄", layout="wide")

if "user" not in st.session_state:
    st.session_state.user = None
if "last_result" not in st.session_state:
    st.session_state.last_result = None

st.title("文档智能处理系统")
st.caption("支持注册登录、用户文件隔离、PDF/Word/图片/扫描版 PDF 解析、OCR、摘要、词频可视化与自动分类。")

if st.session_state.user is None:
    login_tab, register_tab = st.tabs(["登录", "注册"])

    with login_tab:
        with st.form("login_form"):
            login_username = st.text_input("用户名", key="login_username")
            login_password = st.text_input("密码", type="password", key="login_password")
            submitted = st.form_submit_button("登录")
        if submitted:
            user = authenticate(login_username, login_password)
            if user is None:
                st.error("用户名或密码错误。")
            else:
                st.session_state.user = user
                st.session_state.last_result = None
                st.rerun()

    with register_tab:
        with st.form("register_form"):
            register_username = st.text_input("用户名", key="register_username")
            register_password = st.text_input("密码", type="password", key="register_password")
            register_password_2 = st.text_input("确认密码", type="password", key="register_password_2")
            submitted = st.form_submit_button("注册")
        if submitted:
            if register_password != register_password_2:
                st.error("两次输入的密码不一致。")
            else:
                ok, message = create_user(register_username, register_password)
                if ok:
                    st.success(message)
                else:
                    st.error(message)
    st.stop()

user = st.session_state.user

with st.sidebar:
    st.write(f"当前用户：{user['username']}")
    if st.button("退出登录"):
        st.session_state.user = None
        st.session_state.last_result = None
        st.rerun()

    st.divider()
    st.subheader("我的文档")
    documents = list_user_documents(user["id"])
    if not documents:
        st.caption("当前账号暂无上传记录。")
        selected_doc_id = None
    else:
        options = {f"{doc['created_at']}  {doc['filename']}": doc["id"] for doc in documents}
        selected_label = st.selectbox("历史记录", list(options.keys()))
        selected_doc_id = options[selected_label]
        if st.button("查看历史文档"):
            doc = get_user_document(user["id"], selected_doc_id)
            if doc is not None:
                category, category_scores, summary, freq_df = analyze_text(doc["extracted_text"])
                st.session_state.last_result = {
                    "filename": doc["filename"],
                    "category": category,
                    "category_scores": category_scores,
                    "summary": summary,
                    "freq_df": freq_df,
                    "extracted_text": doc["extracted_text"],
                    "from_history": True,
                }
                st.rerun()

uploaded_file = st.file_uploader(
    "上传文档",
    type=["pdf", "docx", "png", "jpg", "jpeg", "bmp", "tif", "tiff"],
)

if uploaded_file is not None:
    file_bytes = uploaded_file.getvalue()
    with st.spinner("正在解析文档，请稍候..."):
        try:
            extracted_text = parse_uploaded_file(uploaded_file)
        except Exception as exc:
            st.error(f"文件解析失败：{exc}")
            st.stop()

    if not extracted_text:
        st.warning("未提取到有效文本。请确认文件内容清晰，扫描件建议使用高分辨率图片或 PDF。")
        st.stop()

    category, category_scores, summary, freq_df = analyze_text(extracted_text)
    save_document(
        user_id=user["id"],
        filename=uploaded_file.name,
        file_bytes=file_bytes,
        category=category,
        summary=summary,
        extracted_text=extracted_text,
        word_count=len(freq_df),
    )
    st.session_state.last_result = {
        "filename": uploaded_file.name,
        "category": category,
        "category_scores": category_scores,
        "summary": summary,
        "freq_df": freq_df,
        "extracted_text": extracted_text,
        "from_history": False,
    }
    st.success("文档已解析并保存到当前用户的历史记录。")

if st.session_state.last_result is None:
    st.info("请上传一个 PDF、Word 或图片文件开始处理。上传记录只会显示在当前登录账号下。")
    st.stop()

result = st.session_state.last_result
st.subheader(f"当前文档：{result['filename']}")
render_analysis(
    result["category"],
    result["category_scores"],
    result["summary"],
    result["freq_df"],
    result["extracted_text"],
)

st.download_button(
    "下载提取文本",
    data=result["extracted_text"].encode("utf-8"),
    file_name=f"{result['filename']}_提取文本.txt",
    mime="text/plain",
)
